from docx import Document
import pandas as pd
import re
from openai import OpenAI
from prompts import get_prompt

def extract_paragraphs(uploaded_file):
    """استخراج الفقرات النصية من ملف وورد"""
    doc = Document(uploaded_file)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return paragraphs

def process_paragraphs_with_gpt(paragraphs, model="gpt-4", api_key=""):
    """تحليل الفقرات باستخدام OpenAI وتحديد الكشافات"""
    client = OpenAI(api_key=api_key)
    results = []

    for i, para in enumerate(paragraphs):
        prompt = get_prompt(para)

        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "أنت مساعد ذكي ومتمكن في فهم وتحليل النصوص الشرعية بدقة."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            reply = response.choices[0].message.content.strip()

            title_lines = []
reason = ""

for line in reply.splitlines():
    line = line.strip()
    if re.match(r"<(/?\d+/.*)>", line):
        title_lines.append(line)
    elif line.startswith("سبب التصنيف:"):
        reason = line.replace("سبب التصنيف:", "").strip()

# حفظ النتائج
for tag in title_lines:
    results.append({
        "الفقرة رقم": i + 1,
        "الكشاف": tag,
        "النص": para,
        "سبب التصنيف": reason
    })

        except Exception as e:
            results.append({
                "الفقرة رقم": i + 1,
                "الكشاف": "خطأ: " + str(e),
                "النص": para
            })

    return results

def insert_titles_to_docx(paragraphs, results):
    """إدراج عناوين الكشاف قبل الفقرات في ملف وورد جديد"""
    doc = Document()
    result_dict = {res["النص"]: res["الكشاف"] for res in results if "الكشاف" in res and res["الكشاف"].startswith("<")}

    for para in paragraphs:
        if para in result_dict:
            doc.add_paragraph(result_dict[para])
        doc.add_paragraph(para)

    return doc
